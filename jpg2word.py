import os
import sys
import threading
from tkinter import filedialog, messagebox
import customtkinter as ctk
import easyocr
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENT, WD_SECTION
from PIL import Image
import numpy as np

# 设置外观
ctk.set_appearance_mode("System")
ctk.set_default_color_theme("blue")

# 环境变量补丁 (防止部分库冲突)
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"

class Jpg2WordApp(ctk.CTk):
    def __init__(self):
        super().__init__()

        self.title("Windows 11 图片转 Word 工具 (稳定版)")
        self.geometry("700x520")

        # 变量
        self.input_paths = []
        self.lang_var = ctk.StringVar(value="ch_sim")  # 默认中文 (EasyOCR: ch_sim)
        self.paper_size_var = ctk.StringVar(value="A4")
        self.use_gpu_var = ctk.BooleanVar(value=False) # 默认 CPU

        self.setup_ui()
        
        # 缓存 Reader 实例
        self.readers = {}

    def setup_ui(self):
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        # 左侧面板
        self.sidebar_frame = ctk.CTkFrame(self, width=170, corner_radius=0)
        self.sidebar_frame.grid(row=0, column=0, rowspan=4, sticky="nsew")
        self.sidebar_frame.grid_rowconfigure(8, weight=1)

        self.logo_label = ctk.CTkLabel(self.sidebar_frame, text="参数设置", font=ctk.CTkFont(size=20, weight="bold"))
        self.logo_label.grid(row=0, column=0, padx=20, pady=(20, 10))

        # 统一组件宽度
        widget_width = 140

        # 语言选择
        self.lang_label = ctk.CTkLabel(self.sidebar_frame, text="识别语言:", anchor="w")
        self.lang_label.grid(row=1, column=0, padx=20, pady=(10, 0))
        self.lang_menu = ctk.CTkOptionMenu(self.sidebar_frame, values=["中文 (简体)", "英文", "印尼语"],
                                         command=self.change_lang_event, width=widget_width)
        self.lang_menu.grid(row=2, column=0, padx=20, pady=(0, 10))

        # 纸张大小
        self.paper_label = ctk.CTkLabel(self.sidebar_frame, text="纸张大小:", anchor="w")
        self.paper_label.grid(row=3, column=0, padx=20, pady=(10, 0))
        self.paper_menu = ctk.CTkOptionMenu(self.sidebar_frame, values=["A4", "A3", "Letter"],
                                          variable=self.paper_size_var, width=widget_width)
        self.paper_menu.grid(row=4, column=0, padx=20, pady=(0, 10))

        # GPU 选项
        self.gpu_switch = ctk.CTkSwitch(self.sidebar_frame, text="使用 GPU", variable=self.use_gpu_var)
        self.gpu_switch.grid(row=5, column=0, padx=20, pady=(20, 10))

        # 主界面
        self.main_frame = ctk.CTkFrame(self, corner_radius=10)
        self.main_frame.grid(row=0, column=1, padx=20, pady=20, sticky="nsew")
        self.main_frame.grid_columnconfigure(0, weight=1)
        self.main_frame.grid_rowconfigure(2, weight=1)

        self.title_label = ctk.CTkLabel(self.main_frame, text="图片转 Word (支持多语言 & 布局保持)", font=ctk.CTkFont(size=18, weight="bold"))
        self.title_label.grid(row=0, column=0, padx=20, pady=(20, 10))

        # 统一按钮大小和间距
        btn_width = 240
        btn_height = 45
        
        self.select_btn = ctk.CTkButton(self.main_frame, text="选择图片文件", command=self.select_images, 
                                       width=btn_width, height=btn_height, font=ctk.CTkFont(size=15))
        self.select_btn.grid(row=1, column=0, padx=20, pady=15)

        self.file_list_box = ctk.CTkTextbox(self.main_frame)
        self.file_list_box.grid(row=2, column=0, padx=20, pady=10, sticky="nsew")
        self.file_list_box.insert("0.0", "未选择文件...\n")

        self.progress_bar = ctk.CTkProgressBar(self.main_frame)
        self.progress_bar.grid(row=3, column=0, padx=20, pady=15, sticky="ew")
        self.progress_bar.set(0)

        self.status_label = ctk.CTkLabel(self.main_frame, text="就绪")
        self.status_label.grid(row=4, column=0, padx=20, pady=5)

        self.convert_btn = ctk.CTkButton(self.main_frame, text="开始转换", command=self.start_conversion, 
                                       fg_color="green", hover_color="darkgreen", 
                                       width=btn_width, height=btn_height,
                                       font=ctk.CTkFont(size=16, weight="bold"))
        self.convert_btn.grid(row=5, column=0, padx=20, pady=(15, 25))

    def change_lang_event(self, new_lang: str):
        if "中文" in new_lang:
            self.lang_var.set("ch_sim")
        elif "英文" in new_lang:
            self.lang_var.set("en")
        elif "印尼语" in new_lang:
            self.lang_var.set("id")

    def select_images(self):
        files = filedialog.askopenfilenames(title="选择图片", filetypes=[("图片文件", "*.jpg *.jpeg *.png *.bmp *.tiff")])
        if files:
            self.input_paths = list(files)
            self.file_list_box.delete("0.0", "end")
            for f in self.input_paths:
                self.file_list_box.insert("end", f + "\n")
            self.status_label.configure(text=f"已选择 {len(self.input_paths)} 个文件")

    def get_reader(self, lang, use_gpu):
        key = (lang, use_gpu)
        if key not in self.readers:
            # EasyOCR 初始化时加载模型
            langs = [lang]
            if lang != 'en':
                langs.append('en') # 默认携带英文支持
            self.readers[key] = easyocr.Reader(langs, gpu=use_gpu)
        return self.readers[key]

    def start_conversion(self):
        if not self.input_paths:
            messagebox.showwarning("警告", "请先选择图片文件！")
            return

        self.convert_btn.configure(state="disabled")
        self.status_label.configure(text="正在初始化识别引擎 (首次运行较慢)...")
        
        threading.Thread(target=self.process_conversion, daemon=True).start()

    def process_conversion(self):
        try:
            # 使用 EasyOCR Reader
            reader = self.get_reader(self.lang_var.get(), self.use_gpu_var.get())
            
            doc = Document()
            # 设置默认字体和样式
            style = doc.styles['Normal']
            font = style.font
            font.name = 'SimSun' # 宋体
            font.size = Pt(10.5) # 五号
            # 中文字体兼容性设置
            style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

            self.set_paper_size(doc)

            total = len(self.input_paths)
            for i, img_path in enumerate(self.input_paths):
                self.status_label.configure(text=f"正在分析第 {i+1}/{total} 张图片...")
                self.progress_bar.set((i) / total)
                
                # 获取图片尺寸用于缩放计算 (如果需要)
                # with Image.open(img_path) as img:
                #    width, height = img.size

                # 识别
                results = reader.readtext(img_path)
                
                if i > 0:
                    doc.add_page_break()

                if results:
                    # 按照 y 坐标排序 (基于框的顶部 y)
                    results.sort(key=lambda x: x[0][0][1])
                    
                    rows = []
                    if results:
                        current_row = [results[0]]
                        for j in range(1, len(results)):
                            # 判断是否为同一行：检查 y 轴上的重叠或距离
                            # 获取前一个块的垂直范围 [y_top, y_bottom]
                            prev_box = current_row[-1][0]
                            prev_y_top = prev_box[0][1]
                            prev_y_bottom = prev_box[2][1]
                            prev_height = prev_y_bottom - prev_y_top

                            curr_box = results[j][0]
                            curr_y_top = curr_box[0][1]
                            curr_y_bottom = curr_box[2][1]

                            # 如果当前块的顶部 y 在前一个块的垂直范围内，或者是重叠度很高，视为同一行
                            # 允许一定误差 (例如高度的 30%)
                            if abs(curr_y_top - prev_y_top) < prev_height * 0.5:
                                current_row.append(results[j])
                            else:
                                rows.append(current_row)
                                current_row = [results[j]]
                        rows.append(current_row)

                    for row in rows:
                        row.sort(key=lambda x: x[0][0][0]) # 每一行按 x 排序
                        p = doc.add_paragraph()
                        # 设置行间距为单倍
                        p.paragraph_format.line_spacing = 1.0
                        p.paragraph_format.space_after = Pt(0)
                        
                        last_x_right = 0
                        for line in row:
                            box, text, _ = line
                            x_left = box[0][0]
                            x_right = box[1][0]
                            box_height = box[2][1] - box[0][1]
                            
                            # 基于文字高度估算单个字符的宽度 (通常高度/2 是一个合理的估算)
                            char_width_est = box_height * 0.5
                            
                            if last_x_right > 0:
                                distance = x_left - last_x_right
                                if distance > char_width_est:
                                    space_count = int(distance / char_width_est)
                                    # 限制最大空格数，防止排版飞散
                                    p.add_run(" " * min(space_count, 30))
                                else:
                                    # 如果距离很小，给 1-2 个空格
                                    p.add_run(" ")
                            
                            p.add_run(text)
                            last_x_right = x_right

            # 自动保存到图片目录
            # 获取第一张图片的目录和文件名作为基础
            first_img_path = self.input_paths[0]
            img_dir = os.path.dirname(first_img_path)
            img_name = os.path.splitext(os.path.basename(first_img_path))[0]
            
            # 构造保存路径 (例如: 图片名_识别结果.docx)
            save_path = os.path.join(img_dir, f"{img_name}_识别结果.docx")
            save_path = os.path.abspath(save_path)
            
            doc.save(save_path)
            # 转换成功，通知主线程处理弹窗和打开操作
            self.after(0, lambda: self.conversion_complete_ui(save_path))

        except Exception as e:
            self.after(0, lambda: messagebox.showerror("错误", f"识别过程中出现意外: {str(e)}"))
            self.after(0, lambda: self.status_label.configure(text="转换失败"))
        finally:
            self.after(0, lambda: self.convert_btn.configure(state="normal"))

    def conversion_complete_ui(self, save_path):
        self.progress_bar.set(1.0)
        self.status_label.configure(text="转换成功！")
        if messagebox.askyesno("完成", f"转换已完成！\n保存在: {save_path}\n是否立即打开？"):
            try:
                # 使用更稳健的方式在 Windows 上打开文件并分离进程
                os.startfile(save_path)
            except Exception as e:
                messagebox.showerror("打开失败", f"无法自动打开文档，请手动查看: {str(e)}")

    def set_paper_size(self, doc):
        section = doc.sections[0]
        # 设置标准页边距 (2.54cm)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        
        size = self.paper_size_var.get()
        if size == "A4":
            section.page_height, section.page_width = Cm(29.7), Cm(21.0)
        elif size == "A3":
            section.page_height, section.page_width = Cm(42.0), Cm(29.7)
        elif size == "Letter":
            section.page_height, section.page_width = Inches(11), Inches(8.5)

if __name__ == "__main__":
    app = Jpg2WordApp()
    app.mainloop()
